from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Any

import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Emu, Inches, Pt, RGBColor
from PIL import Image as PILImage

from app.core.logging_config import get_logger

logger = get_logger(__name__)

_RTL_FONTS = ["Traditional Arabic", "Arial", "Times New Roman", "Calibri"]
_DEFAULT_FONT = "Arial"
_FONT_SIZES = {
    "heading1": 22,
    "heading2": 18,
    "heading3": 16,
    "heading4": 14,
    "body": 12,
    "small": 10,
}


def _reshape_arabic(text: str) -> str:
    try:
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    except Exception:
        return text


def _has_arabic(text: str) -> bool:
    arabic_pattern = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]")
    return bool(arabic_pattern.search(text))


def _get_paragraph_style(size_pt: int, bold: bool = False, italic: bool = False) -> dict:
    return {
        "size": Pt(size_pt),
        "bold": bold,
        "italic": italic,
    }


def set_rtl_direction(document: Document) -> None:
    for section in document.sections:
        sect_pr = section._sectPr
        if sect_pr is None:
            sect_pr = parse_xml(f'<w:sectPr {nsdecls("w")}></w:sectPr>')
            section._sectPr = sect_pr

        bidi = sect_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
            sect_pr.append(bidi)
        else:
            bidi.set(qn("w:val"), "1")

        text_dir = sect_pr.find(qn("w:textDirection"))
        if text_dir is None:
            text_dir = parse_xml(f'<w:textDirection {nsdecls("w")} w:val="lrTb"/>')
            sect_pr.append(text_dir)
        else:
            text_dir.set(qn("w:val"), "lrTb")


def set_margins(
    document: Document,
    top: float = 1.0,
    bottom: float = 1.0,
    left: float = 1.0,
    right: float = 1.0,
) -> None:
    for section in document.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def set_fonts(document: Document, font_name: str = _DEFAULT_FONT) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = font_name
    rpr = style.element.rPr
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:cs"), font_name)


def _set_run_font(run, font_name: str = _DEFAULT_FONT, size: int = 12, bold: bool = False, italic: bool = False,
                  color: str | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color.lstrip("#")))

    rpr = run._element.rPr
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._element.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:cs"), font_name)

    if "Traditional Arabic" in font_name:
        sz = rpr.find(qn("w:sz"))
        if sz is None:
            sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{size * 2}"/>')
            rpr.append(sz)


def add_paragraphs(
    document: Document,
    text: str,
    font_name: str = _DEFAULT_FONT,
    font_size: int = 12,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    if not text.strip():
        return

    has_arabic = _has_arabic(text)
    display_text = _reshape_arabic(text) if has_arabic else text

    paragraph = document.add_paragraph()
    run = paragraph.add_run(display_text)

    _set_run_font(run, font_name=font_name, size=font_size, bold=bold, italic=italic, color=color)

    if alignment is not None:
        paragraph.alignment = alignment
    elif has_arabic:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    pPr = paragraph._element.pPr
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        paragraph._element.insert(0, pPr)

    if has_arabic:
        bidi = pPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
            pPr.append(bidi)
        else:
            bidi.set(qn("w:val"), "1")


def add_heading(
    document: Document,
    text: str,
    level: int = 1,
    font_name: str = _DEFAULT_FONT,
) -> None:
    if not text.strip():
        return

    has_arabic = _has_arabic(text)
    display_text = _reshape_arabic(text) if has_arabic else text

    heading = document.add_heading(display_text, level=min(level, 4))

    for run in heading.runs:
        _set_run_font(run, font_name=font_name)

    if has_arabic:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    pPr = heading._element.pPr
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        heading._element.insert(0, pPr)

    if has_arabic:
        bidi = pPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
            pPr.append(bidi)


def add_tables(
    document: Document,
    tables_data: list[list[list[str]]],
    font_name: str = _DEFAULT_FONT,
) -> None:
    for table_data in tables_data:
        if not table_data:
            continue

        rows_count = len(table_data)
        cols_count = max(len(row) for row in table_data) if table_data else 0

        if cols_count == 0:
            continue

        table = document.add_table(rows=rows_count, cols=cols_count)
        table.style = "Table Grid"

        for row_idx, row_data in enumerate(table_data):
            for col_idx in range(cols_count):
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                cell = table.cell(row_idx, col_idx)
                cell.text = ""

                has_arabic = _has_arabic(cell_text)
                display_text = _reshape_arabic(cell_text) if has_arabic else cell_text

                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(display_text)
                _set_run_font(run, font_name=font_name, size=10)

                if has_arabic:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph()


def add_images(
    document: Document,
    images: list[dict[str, Any]],
    max_width: int = 400,
) -> None:
    for img_data in images:
        pil_image: PILImage.Image = img_data.get("image")
        if pil_image is None:
            continue

        img_width, img_height = pil_image.size
        aspect = img_height / img_width if img_width > 0 else 1

        if img_width > max_width:
            img_width = max_width
            img_height = int(img_width * aspect)

        img_bytes = io.BytesIO()
        pil_image.save(img_bytes, format="PNG")
        img_bytes.seek(0)

        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_bytes, width=Inches(img_width / 96))

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def preserve_formatting(
    document: Document,
    text_blocks: list[dict[str, Any]],
    font_name: str = _DEFAULT_FONT,
) -> None:
    for block in text_blocks:
        block_type = block.get("type", "paragraph")
        text = block.get("text", "")
        if not text.strip():
            continue

        if block_type == "heading":
            add_heading(
                document,
                text,
                level=block.get("level", 1),
                font_name=font_name,
            )
        elif block_type == "image":
            images = block.get("images", [])
            add_images(document, images)
        elif block_type == "table":
            add_tables(document, block.get("tables", []), font_name=font_name)
        else:
            add_paragraphs(
                document,
                text,
                font_name=font_name,
                font_size=block.get("font_size", 12),
                bold=block.get("bold", False),
                italic=block.get("italic", False),
                color=block.get("color"),
            )


def create_docx(
    text_content: str,
    output_path: str | Path,
    font_name: str = _DEFAULT_FONT,
    title: str = "Converted Document",
    text_blocks: list[dict[str, Any]] | None = None,
    images: list[dict[str, Any]] | None = None,
    tables_data: list[list[list[str]]] | None = None,
    margins: tuple[float, float, float, float] | None = None,
) -> Path:
    document = Document()

    set_rtl_direction(document)
    set_fonts(document, font_name=font_name)

    if margins:
        set_margins(document, *margins)
    else:
        set_margins(document)

    core_properties = document.core_properties
    core_properties.title = title

    if text_blocks:
        preserve_formatting(document, text_blocks, font_name=font_name)
    else:
        for paragraph_text in text_content.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue

            if paragraph_text.startswith("--- Page") and paragraph_text.endswith("---"):
                add_paragraphs(
                    document,
                    paragraph_text,
                    font_name=font_name,
                    font_size=_FONT_SIZES["body"],
                    bold=True,
                )
                continue

            add_paragraphs(
                document,
                paragraph_text,
                font_name=font_name,
                font_size=_FONT_SIZES["body"],
            )

    if tables_data:
        add_tables(document, tables_data, font_name=font_name)

    if images:
        document.add_paragraph()
        add_images(document, images)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document.save(str(output_path))
    logger.info("DOCX saved to %s", output_path)

    return output_path
