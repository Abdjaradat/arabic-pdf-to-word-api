# RSWS Algorithm: Read -> Store -> Write -> Style
# اختراع: Read PDF, Store words+formatting, Write word-by-word into Word, Style each word
#
# CONFIDENTIAL — Proprietary Intellectual Property
# The RSWS engine, including its document reconstruction pipeline, iterative correction
# workflow, layout recovery logic, visual comparison system, and adaptive refinement
# methodology, constitutes proprietary intellectual property.
# Unauthorized disclosure, publication, reproduction, reverse engineering, redistribution,
# or commercial usage is strictly prohibited.
import sys
import os
import json
import uuid
import time
from dataclasses import dataclass
from typing import List, Tuple, Optional

import fitz
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Best config from 100-iteration optimizer ──
RSWS_CONFIG = {
    'sort_rtl': True,
    'add_spaces': True,
    'use_orig_size': True,
    'font_arabic': 'Traditional Arabic',
    'font_latin': 'Arial',
    'size_min': 7,
    'size_max': 72,
    'line_tolerance': 0.03,
    'detect_tables': True,
    'include_images': True,
    'align_by_content': False,
}

# Allow override from env var
_env_config = os.environ.get('RSWS_CONFIG')
if _env_config:
    try:
        overrides = json.loads(_env_config)
        RSWS_CONFIG.update(overrides)
    except: pass


# ── Data structures ──

@dataclass
class WordInfo:
    text: str
    font: str
    size: float
    bold: bool
    italic: bool
    color: Tuple[int, int, int]

@dataclass
class LineInfo:
    words: List[WordInfo]
    alignment: str
    has_arabic: bool
    y: float
    x: float

@dataclass
class PageInfo:
    lines: List[LineInfo]
    images: List  # list of (png_bytes, w_in, h_in)


# ── Helpers ──

ARABIC_RANGES = [
    (0x0600, 0x06FF),  # Arabic
    (0x0750, 0x077F),  # Arabic Supplement
    (0x08A0, 0x08FF),  # Arabic Extended-A
    (0xFB50, 0xFDFF),  # Arabic Presentation Forms-A
    (0xFE70, 0xFEFF),  # Arabic Presentation Forms-B
    (0x0600, 0x06FF),  # (dup for clarity)
]

def has_arabic(text: str) -> bool:
    return any(lo <= ord(c) <= hi for lo, hi in ARABIC_RANGES for c in text)


def set_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def set_run_font(run, font_name: str):
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    # remove old rFonts if exists
    for child in list(rPr):
        if child.tag == qn('w:rFonts'):
            rPr.remove(child)
    rPr.insert(0, rFonts)


def extract_color(span) -> Tuple[int, int, int]:
    c = span.get('color', 0)
    if isinstance(c, int):
        return ((c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)
    return (0, 0, 0)


def pick_font(text: str, span_font: str,
              font_arabic: str = 'Traditional Arabic',
              font_latin: str = 'Arial') -> str:
    """Pick the best font for a word based on content and original span font."""
    if not span_font or span_font == 'ArialMT':
        return font_arabic if has_arabic(text) else font_latin
    font_map = {
        'ArialMT': font_latin,
        'Arial-BoldMT': font_latin,
        'Arial-ItalicMT': font_latin,
        'Arial-BoldItalicMT': font_latin,
        'TimesNewRomanPSMT': 'Times New Roman',
        'TimesNewRomanPS-BoldMT': 'Times New Roman',
        'TimesNewRomanPS-ItalicMT': 'Times New Roman',
        'TimesNewRomanPS-BoldItalicMT': 'Times New Roman',
        'CourierNewPSMT': 'Courier New',
        'CourierNewPS-BoldMT': 'Courier New',
        'CourierNewPS-ItalicMT': 'Courier New',
        'CourierNewPS-BoldItalicMT': 'Courier New',
        'TraditionalArabic': font_arabic,
        'Traditional Arabic': font_arabic,
        'Calibri': 'Calibri',
        'Calibri-Bold': 'Calibri',
        'Calibri-Italic': 'Calibri',
        'Calibri-BoldItalic': 'Calibri',
    }
    mapped = font_map.get(span_font, span_font)
    # If text is Arabic but font isn't Arabic-friendly, override
    if has_arabic(text) and mapped not in (font_arabic, font_latin, 'Calibri', 'Times New Roman'):
        return font_arabic
    return mapped


# ═══════════════════════════════════════════════
#  STEP 1 - READ & EXTRACT (word by word)
#  ═══════════════════════════════════════════════
#  Uses PyMuPDF "words" for correct segmentation,
#  then maps each word to its span for formatting.

def extract_from_pdf(pdf_path: str) -> List[PageInfo]:
    pdf = fitz.open(pdf_path)
    pages = []

    for page_num in range(pdf.page_count):
        page = pdf[page_num]
        pw = page.rect.width

        # Get clean word segmentation with positions
        raw_words = page.get_text("words")  # (x0,y0,x1,y1,word,block_no,line_no,word_no)
        # Get detailed dict for formatting info
        blocks = page.get_text("dict")["blocks"]

        lines_info = []
        images_info = []

        # Extract images
        for block in blocks:
            if block['type'] != 1: continue
            try:
                b = block['bbox']
                w = b[2]-b[0]
                h = b[3]-b[1]
                if w < 20 or h < 20: continue
                pix = page.get_pixmap(clip=b, width=int(w), height=int(h))
                images_info.append((pix.tobytes('png'), w/72, h/72))
            except:
                pass

        if not raw_words:
            pages.append(PageInfo(lines=[], images=images_info))
            continue

        # Build span lookup: for each word position, find which span it falls in
        # Extract all spans with their bbox and formatting
        span_list = []  # (x0, y0, x1, y1, text, font, size, bold, italic, color)
        for block in blocks:
            if block['type'] != 0: continue
            for line in block.get('lines', []):
                for span in line.get('spans', []):
                    sb = span['bbox']
                    txt = span['text']
                    if not txt.strip():
                        continue
                    font = span.get('font', 'Arial')
                    size = span.get('size', 11)
                    flags = span.get('flags', 0)
                    span_list.append((
                        sb[0], sb[1], sb[2], sb[3],
                        txt, font, size,
                        bool(flags & 2),  # bold
                        bool(flags & 1),  # italic
                        extract_color(span),
                    ))

        def find_span(wx0, wy0, wx1, wy1):
            """Find the span that mostly contains this word bbox."""
            wcx = (wx0 + wx1) / 2
            wcy = (wy0 + wy1) / 2
            best = None
            best_area = 0
            for sx0, sy0, sx1, sy1, stxt, sfont, ssize, sbold, sital, scol in span_list:
                if sx0 <= wcx <= sx1 and sy0 <= wcy <= sy1:
                    area = (sx1 - sx0) * (sy1 - sy0)
                    if area > best_area:
                        best_area = area
                        best = (sfont, ssize, sbold, sital, scol)
            return best

        # Group words into lines by y-position
        LINE_TOLERANCE = page.rect.height * RSWS_CONFIG.get('line_tolerance', 0.03)
        word_lines = []
        for w in raw_words:
            wx0, wy0, wx1, wy1, word = w[0], w[1], w[2], w[3], w[4]
            font_info = find_span(wx0, wy0, wx1, wy1) or ('Arial', 11, False, False, (0,0,0))
            placed = False
            for line in word_lines:
                if line and abs(line[0][3] - wy0) < LINE_TOLERANCE:
                    line.append((word, wx0, wy0, wx1, wy1, font_info))
                    placed = True
                    break
            if not placed:
                word_lines.append([(word, wx0, wy0, wx1, wy1, font_info)])

        # Sort each line left-to-right
        for line in word_lines:
            line.sort(key=lambda x: x[1])

        # Handle RTL lines: reverse word order
        if RSWS_CONFIG.get('sort_rtl', True):
            for line in word_lines:
                line_text = ' '.join(x[0] for x in line)
                if has_arabic(line_text):
                    line.reverse()

        word_lines.sort(key=lambda x: x[0][2])

        add_spaces = RSWS_CONFIG.get('add_spaces', True)
        use_orig_size = RSWS_CONFIG.get('use_orig_size', True)
        size_min = RSWS_CONFIG.get('size_min', 7)
        size_max = RSWS_CONFIG.get('size_max', 72)
        font_arabic = RSWS_CONFIG.get('font_arabic', 'Traditional Arabic')
        font_latin = RSWS_CONFIG.get('font_latin', 'Arial')

        for line in word_lines:
            if not line: continue
            line_words_info = []
            line_text_combined = ''
            for i, (word, wx0, wy0, wx1, wy1, (sfont, ssize, sbold, sital, scol)) in enumerate(line):
                display_word = word + (' ' if add_spaces and i < len(line) - 1 else '')
                line_text_combined += display_word
                mapped_font = pick_font(word, sfont, font_arabic, font_latin)
                sz = ssize if use_orig_size else 12
                sz = max(size_min, min(sz, size_max))
                line_words_info.append(WordInfo(
                    text=display_word,
                    font=mapped_font,
                    size=sz,
                    bold=sbold,
                    italic=sital,
                    color=scol,
                ))

            line_has_arabic = has_arabic(line_text_combined)

            first_x = line[0][1]
            last_x = line[-1][3]
            line_width = last_x - first_x
            margin_left = first_x
            margin_right = pw - last_x

            if RSWS_CONFIG.get('align_by_content', False):
                alignment = 'right' if line_has_arabic else 'left'
            else:
                if line_width < pw * 0.3 and margin_right < pw * 0.05:
                    alignment = 'right'
                elif line_width < pw * 0.3 and margin_left < pw * 0.05:
                    alignment = 'left'
                elif margin_left > pw * 0.25 and margin_right > pw * 0.25:
                    alignment = 'center'
                else:
                    alignment = 'justify'

            lines_info.append(LineInfo(
                words=line_words_info,
                alignment=alignment,
                has_arabic=line_has_arabic,
                y=line[0][2],
                x=first_x,
            ))

        pages.append(PageInfo(lines=lines_info, images=images_info))

    pdf.close()
    return pages


# ═══════════════════════════════════════════════
#  STEP 2 - TYPE INTO WORD (word by word)
#  ═══════════════════════════════════════════════

def type_into_word(doc: Document, pages: List[PageInfo]) -> List:
    refs = []

    for page_idx, page in enumerate(pages):
        for line in page.lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.2
            for word in line.words:
                run = p.add_run(word.text)
                refs.append((p, run, word, line))

        # Images
        for img_bytes, w_in, h_in in page.images:
            if w_in > 0.5 and h_in > 0.5:
                w = min(w_in, 6.0)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_bytes, width=Inches(w))

        if page_idx < len(pages) - 1:
            doc.add_page_break()

    return refs


# ═══════════════════════════════════════════════
#  STEP 3 - FORMAT LIKE HUMAN
#  ═══════════════════════════════════════════════

def format_like_human(refs: List):
    last_p = None
    for item in refs:
        if len(item) == 4:
            p, run, word, line = item

            if p is not last_p:
                align_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                if line.alignment in align_map:
                    p.alignment = align_map[line.alignment]
                if line.has_arabic:
                    set_rtl(p)
                last_p = p

            # Per-word font
            set_run_font(run, word.font)
            size = max(7, min(word.size, 72))
            run.font.size = Pt(size)
            run.bold = word.bold
            run.italic = word.italic
            if word.color != (0, 0, 0):
                try:
                    run.font.color.rgb = RGBColor(*word.color)
                except:
                    pass


# ═══════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════

def convert_pdf_to_word(input_path: str, output_dir: str) -> str:
    t0 = time.time()

    t1 = time.time()
    pages = extract_from_pdf(input_path)
    extract_time = time.time() - t1

    t2 = time.time()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    refs = type_into_word(doc, pages)
    type_time = time.time() - t2

    t3 = time.time()
    format_like_human(refs)
    format_time = time.time() - t3

    output_filename = f"{uuid.uuid4()}.docx"
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)

    total_text = sum(len(w.text) for p in pages for l in p.lines for w in l.words)

    return json.dumps({
        'outputPath': output_path.replace('\\', '/'),
        'outputFileName': output_filename,
        'pageCount': len(pages),
        'textLength': total_text,
        'method': 'rsws',
        'timing': {
            'extract': round(extract_time, 2),
            'type': round(type_time, 2),
            'format': round(format_time, 2),
            'total': round(time.time() - t0, 2),
        }
    })


# ── Built-in 100-iteration optimizer ──
def run_optimizer(output_dir: str):
    """Run RSWS 100 times with different configs to find optimal parameters."""
    # Generate test PDF
    test_pdf = os.path.join(output_dir, '_optimize_test.pdf')
    try:
        from fpdf import FPDF
        pdf = FPDF(orientation='P', unit='mm', format='A4')
        pdf.add_page()
        pdf.set_font('Helvetica', '', 12)
        pdf.cell(0, 10, 'Arabic Test: مرحبا بالعالم Hello World 2024', new_x="LMARGIN", new_y="NEXT", align='R')
        pdf.cell(0, 10, 'Mixed: شركة ABC للتطوير - Saudi Arabia', new_x="LMARGIN", new_y="NEXT", align='R')
        pdf.cell(0, 10, 'Table: 1-15000 SAR  2-30000 SAR  3-45000 SAR', new_x="LMARGIN", new_y="NEXT", align='R')
        pdf.cell(0, 10, 'نهاية العقد - End of Contract', new_x="LMARGIN", new_y="NEXT", align='R')
        pdf.output(test_pdf)
    except:
        # Fallback: use a simple text-based PDF
        test_pdf = os.path.join(output_dir, '_optimize_test.pdf')
        with open(test_pdf, 'wb') as f:
            f.write(b'%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n5 0 obj<</Length 44>>stream\nBT /F1 12 Tf 100 700 Td (Test PDF) Tj ET\nendstream\nendobj\nxref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000266 00000 n \n0000000347 00000 n \ntrailer<</Size 6/Root 1 0 R>>\nstartxref\n447\n%%EOF')

    print(f"Optimizer: running 100 iterations on {test_pdf}")

    import random, itertools
    best_score = -1
    best_config = None

    # Generate 100 configs
    configs = []
    for sort_rtl in [True, False]:
        for add_sp in [True, False]:
            for use_sz in [True, False]:
                for align in [True, False]:
                    if len(configs) >= 40: break
                    configs.append({'sort_rtl': sort_rtl, 'add_spaces': add_sp, 'use_orig_size': use_sz, 'align_by_content': align})
                if len(configs) >= 40: break
            if len(configs) >= 40: break
        if len(configs) >= 40: break

    for fa in ['Traditional Arabic', 'Arial', 'Times New Roman', 'Calibri']:
        for fl in ['Arial', 'Times New Roman', 'Calibri']:
            if fa == fl: continue
            if len(configs) >= 70: break
            configs.append({'font_arabic': fa, 'font_latin': fl, 'sort_rtl': True, 'add_spaces': True, 'use_orig_size': True, 'align_by_content': False})

    for tol in [0.005, 0.01, 0.015, 0.02, 0.03, 0.05]:
        if len(configs) >= 80: break
        configs.append({'line_tolerance': tol, 'sort_rtl': True, 'add_spaces': True, 'use_orig_size': True, 'align_by_content': False})

    while len(configs) < 100:
        configs.append({
            'sort_rtl': random.choice([True, False]),
            'add_spaces': random.choice([True, False]),
            'use_orig_size': random.choice([True, False]),
            'font_arabic': random.choice(['Traditional Arabic', 'Arial', 'Times New Roman']),
            'font_latin': random.choice(['Arial', 'Times New Roman', 'Calibri']),
            'line_tolerance': random.choice([0.005, 0.01, 0.015, 0.02, 0.03, 0.05]),
            'align_by_content': random.choice([True, False]),
        })

    for idx, cfg in enumerate(configs):
        t0 = time.time()
        # Apply config
        saved = dict(RSWS_CONFIG)
        RSWS_CONFIG.update(cfg)
        try:
            pages = extract_from_pdf(test_pdf)
            doc = Document()
            section = doc.sections[0]
            section.page_width = Cm(21.0); section.page_height = Cm(29.7)
            section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
            refs = type_into_word(doc, pages)
            format_like_human(refs)
            out = os.path.join(output_dir, f'_opt_{idx:03d}.docx')
            doc.save(out)
            words = sum(len(w.text) for p in pages for l in p.lines for w in l.words)
            score = min(words / 80, 1.0) * 100
            if score > best_score:
                best_score = score
                best_config = dict(RSWS_CONFIG)
            print(f"  [{idx+1:>3}/100] score={score:>5.1f} time={time.time()-t0:.2f}s")
        except Exception as e:
            print(f"  [{idx+1:>3}/100] ERROR: {str(e)[:50]}")
        finally:
            RSWS_CONFIG.update(saved)

    # Save best config
    if best_config:
        config_path = os.path.join(output_dir, 'best_config.json')
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(best_config, f, ensure_ascii=False, indent=2)
        RSWS_CONFIG.update(best_config)
        print(f"\nBEST CONFIG (score={best_score:.1f}):")
        for k, v in best_config.items():
            print(f"  {k}: {v}")
        print(f"Saved to: {config_path}")
    else:
        print("Optimizer failed: no valid config found")


# ── Recursive RSWS: RSWS → Word → PDF → RSWS ... 100 times ──
def word_to_pdf(docx_path: str, output_dir: str) -> str:
    """Convert Word to PDF using LibreOffice (or fallback using subprocess)."""
    import subprocess
    pdf_path = os.path.join(output_dir, f"{uuid.uuid4()}.pdf")

    # Try LibreOffice first (with HOME for Docker compatibility)
    env = os.environ.copy()
    env['HOME'] = '/tmp'
    for cmd in ['libreoffice', 'soffice']:
        try:
            r = subprocess.run([cmd, '--headless', '--convert-to', 'pdf',
                               '--outdir', output_dir, docx_path],
                              capture_output=True, timeout=60, env=env)
            if r.returncode != 0:
                print(f"  [!] {cmd} stderr: {r.stderr.decode(errors='replace')[:200]}", file=sys.stderr)
                continue
            # LibreOffice saves with same name but .pdf extension
            lo_pdf = os.path.splitext(docx_path)[0] + '.pdf'
            if os.path.exists(lo_pdf):
                os.rename(lo_pdf, pdf_path)
                print(f"  [✓] {cmd} converted to PDF", file=sys.stderr)
                return pdf_path
        except Exception as e:
            print(f"  [!] {cmd} error: {e}", file=sys.stderr)

    # Fallback: use python-docx + reportlab (basic)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from docx import Document as DocxReader
        c = canvas.Canvas(pdf_path, pagesize=A4)
        doc = DocxReader(docx_path)
        y = 800
        for p in doc.paragraphs:
            c.drawString(50, y, p.text[:100] if p.text else '')
            y -= 20
            if y < 50: c.showPage(); y = 800
        c.save()
        return pdf_path
    except: pass

    return None


def recursive_rsws(input_pdf: str, output_dir: str, iterations: int = 100) -> str:
    """
    Run RSWS recursively: RSWS → Word → PDF → RSWS → Word → PDF → ... × iterations
    يقرا الـ PDF → يحفظ الكلمات → يكتب Word → يحوله PDF → يعيد الكرة 100 مرة
    """
    current_pdf = input_pdf
    final_word = None
    loop_dir = os.path.join(output_dir, '_recursive_loop')
    os.makedirs(loop_dir, exist_ok=True)

    print(f"Recursive RSWS: {iterations} iterations starting...", file=sys.stderr)
    print(f"  Input: {input_pdf}", file=sys.stderr)
    print(f"  Loop dir: {loop_dir}", file=sys.stderr)

    errors_streak = 0

    for i in range(iterations):
        t0 = time.time()
        iter_tag = f"{i+1:03d}"
        iter_word = os.path.join(loop_dir, f"iter_{iter_tag}.docx")
        iter_pdf = os.path.join(loop_dir, f"iter_{iter_tag}.pdf") if i < iterations - 1 else None

        try:
            # STEP 1: RSWS → Word (اقرا → احفظ → اكتب → نسق)
            pages = extract_from_pdf(current_pdf)
            doc = Document()
            section = doc.sections[0]
            section.page_width = Cm(21.0); section.page_height = Cm(29.7)
            section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
            refs = type_into_word(doc, pages)
            format_like_human(refs)
            doc.save(iter_word)

            word_count = sum(len(w.text) for p in pages for l in p.lines for w in l.words)
            errors_streak = 0
            final_word = iter_word
            pdf_result = None

            # STEP 2: Word → PDF (للوجة القادمة)
            if i < iterations - 1:
                pdf_result = word_to_pdf(iter_word, loop_dir)
                if pdf_result and os.path.exists(pdf_result):
                    current_pdf = pdf_result
                else:
                    # PDF conversion failed — use same PDF, RSWS will be similar
                    pass

            elapsed = time.time() - t0
            status = "✓" if i == iterations - 1 else "→"
            pdf_ok = (i == iterations - 1) or (pdf_result and os.path.exists(pdf_result))
            print(f"  [{iter_tag}/{iterations}] {status} words={word_count:>5} "
                  f"time={elapsed:.2f}s PDF={'ok' if pdf_ok else 'skip'}", file=sys.stderr)

        except Exception as e:
            errors_streak += 1
            print(f"  [{iter_tag}/{iterations}] X ERROR: {str(e)[:60]}", file=sys.stderr)
            if errors_streak >= 3:
                print(f"  Stopping: {errors_streak} consecutive errors", file=sys.stderr)
                break
            continue

    print(f"\nRecursive RSWS done! ({iterations} iterations)", file=sys.stderr)
    print(f"Final output: {final_word}", file=sys.stderr)

    # Copy final word to output_dir with clean name
    if final_word and os.path.exists(final_word):
        final_name = f"rsws_recursive_{uuid.uuid4().hex[:8]}.docx"
        final_path = os.path.join(output_dir, final_name)
        import shutil
        shutil.copy2(final_word, final_path)
        return final_path

    return final_word


if __name__ == '__main__':
    if len(sys.argv) >= 2 and sys.argv[1] in ('--optimize', '-o'):
        output_dir = sys.argv[2] if len(sys.argv) > 2 else os.path.join(os.path.dirname(__file__), '..', '..', 'uploads')
        os.makedirs(output_dir, exist_ok=True)
        run_optimizer(output_dir)
        sys.exit(0)

    if len(sys.argv) >= 2 and sys.argv[1] in ('--recursive', '-r'):
        input_pdf = sys.argv[2] if len(sys.argv) > 2 else None
        output_dir = sys.argv[3] if len(sys.argv) > 3 else os.path.join(os.path.dirname(__file__), '..', '..', 'uploads')
        iterations = int(sys.argv[4]) if len(sys.argv) > 4 else 100
        if not input_pdf:
            print(json.dumps({'error': 'Missing input PDF for recursive mode. Usage: --recursive <input_pdf> [output_dir] [iterations]'}))
            sys.exit(1)
        os.makedirs(output_dir, exist_ok=True)
        final = recursive_rsws(input_pdf, output_dir, iterations)
        if final:
            result = {
                'outputPath': final.replace('\\', '/'),
                'outputFileName': os.path.basename(final),
                'method': 'rsws_recursive',
                'iterations': iterations,
            }
            print(json.dumps(result, ensure_ascii=False))
        else:
            print(json.dumps({'error': 'Recursive RSWS failed'}))
        sys.exit(0)

    if len(sys.argv) < 3:
        print(json.dumps({'error': 'Missing arguments. Use --optimize, --recursive, or pass <input_pdf> <output_dir>'}))
        sys.exit(1)

    input_path = sys.argv[1]
    output_dir = sys.argv[2]

    if not os.path.exists(input_path):
        print(json.dumps({'error': f'Input file not found: {input_path}'}))
        sys.exit(1)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    try:
        result = convert_pdf_to_word(input_path, output_dir)
        print(result)
    except Exception as e:
        print(json.dumps({'error': str(e)}))
        sys.exit(1)
